import random
import datetime
import os

import xml.etree.ElementTree as Et
import zipfile

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from datas import (
    main_peoples,
    peoples,
)


def sorted_patrol_list_past(list_1: list, list_2: list) -> list:
    temp_list_1 = []
    temp_list_2 = []

    for x in list_1:
        if list_1.index(x) < 2:
            temp_list_1.append(x)
        else:
            temp_list_2.append(x)

    for x in list_2:
        if list_2.index(x) < 2:
            temp_list_1.append(x)
        else:
            temp_list_2.append(x)
    return temp_list_1 + temp_list_2


def iters_to_docx_tb(cnt: int, rec: list, peps: list, to_stop: int, xml_file: list):
    print("Primary list")
    print(peps)
    random.shuffle(peps)
    for x in peps:
        if x not in xml_file:
            rec.append(
                (
                    peps.index(x) + cnt,
                    x[0],
                    x[1],
                    x[2]
                )
            )

        if len(rec) == to_stop:
            break


def extract_xml(zp_docx: zipfile.ZipFile, date_day, zp_docx_2=None, date_before_yesterday=None):
    try:
        os.mkdir(f"./word_{date_day}")
    except FileExistsError:
        print("File exists. I do not create the file!")
    zp_docx.namelist()
    zp_docx.extractall(f"./word_{date_day}")
    cur_xml = Et.parse(f"./word_{date_day}/word/document.xml")
    cur_xml_read = cur_xml.getroot()
    not_patrol = parse_xml(cur_xml_read)
    if date_before_yesterday is None and zp_docx_2 is None:
        print("Not_patrol")
        print(not_patrol)
        docs(not_patrol)
    else:
        zp_docx_2.namelist()
        zp_docx_2.extractall(f"./word_{date_before_yesterday}")
        before_cur_xml = Et.parse(f"./word_{date_before_yesterday}/word/document.xml")
        before_cur_xml_read = before_cur_xml.getroot()
        before_not_patrol = parse_xml(before_cur_xml_read)
        print("Sum_not_patrol")
        print(sorted_patrol_list_past(not_patrol, before_not_patrol))
        docs(
            sorted_patrol_list_past(not_patrol, before_not_patrol)
        )


def parse_xml(root_of_xml) -> list:
    temp = []
    for x in range(13):
        temp_inner = []
        for y in range(3):
            temp_inner.append(root_of_xml[0][1][x + 3][y + 1][1][0][0].text.strip())
        temp_inner = tuple(temp_inner)
        temp.append(temp_inner)

    return temp


def get_data_from_docx():
    yesterday = datetime.date.today() - datetime.timedelta(days=1)
    day_before_yesterday = datetime.date.today() - datetime.timedelta(days=2)

    if os.path.exists(f"./Patrol_{day_before_yesterday}.docx") and os.path.exists(f"./Patrol_{yesterday}.docx"):
        with zipfile.ZipFile(
                f"Patrol_{yesterday}.docx"
        ) as zp_docx, zipfile.ZipFile(
            f"Patrol_{day_before_yesterday}.docx"
        ) as zp_docx_2:
            extract_xml(zp_docx, yesterday, zp_docx_2, day_before_yesterday)
        print("IF")
    elif os.path.exists(f"./Patrol_{day_before_yesterday}.docx"):
        with zipfile.ZipFile(f"Patrol_{day_before_yesterday}.docx") as zp_docx:
            extract_xml(zp_docx, day_before_yesterday)
        print("ELIF 1")
    elif os.path.exists(f"./Patrol_{yesterday}.docx"):
        with zipfile.ZipFile(f"Patrol_{yesterday}.docx") as zp_docx:
            extract_xml(zp_docx, yesterday)
        print("ELIF 2")
    else:
        print("ELSE")
        docs([])


def docs(xml_file: list):
    document = Document()

    d = document.add_heading(f'Patrol tomorrow {datetime.date.today()}', 0)
    d.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    records = []
    iters_to_docx_tb(1, records, main_peoples, 2, xml_file)
    iters_to_docx_tb(3, records, peoples, 14, xml_file)
    print(records)

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Job'
    hdr_cells[2].text = 'Name'
    hdr_cells[3].text = 'NumberAK'
    for num, qty, ids, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = str(qty)
        row_cells[2].text = ids
        row_cells[3].text = desc

    document.add_page_break()

    document.save(f'Patrol_{datetime.datetime.today().date()}.docx')


get_data_from_docx()
